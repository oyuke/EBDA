from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any
from datetime import datetime
import io

class ReportGenerator:
    def __init__(self, config: Any):
        self.config = config

    def generate_docx(self, wave_data: Dict[str, Any], decision_states: List[Any], snapshot_id: str) -> io.BytesIO:
        """
        Generates a DOCX report summarizing the decision wave.
        Returns bytes buffer.
        """
        doc = Document()
        
        # 1. Title
        title = doc.add_heading(f"Decision Memo: {self.config.customer_name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        meta = doc.add_paragraph()
        meta.add_run(f"Snapshot ID: {snapshot_id}\n").bold = True
        meta.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        meta.add_run(f"Status: {wave_data.get('status', 'DRAFT')}")
        meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph("Based on the evidence collected, the following decisions are recommended for review.")

        # Summary Table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Decision Topic'
        hdr_cells[2].text = 'Status'
        hdr_cells[3].text = 'Recommendation'

        for state in decision_states:
            # state is tuple (card_config, state_obj, score_res)
            card, s, score = state
            
            row_cells = table.add_row().cells
            row_cells[0].text = card.id
            row_cells[1].text = card.title
            row_cells[2].text = s.status
            
            rec_text = "N/A"
            if s.recommendation_draft:
                rec_text = s.recommendation_draft.action
            if s.human_decision_status:
                rec_text = f"[{s.human_decision_status}] {rec_text}"
            
            row_cells[3].text = rec_text

        # 2. Detailed Cards
        doc.add_page_break()
        doc.add_heading("Detailed Analysis", level=1)

        for state in decision_states:
            card, s, score = state
            
            # Card Header
            h = doc.add_heading(f"{card.id}: {card.title}", level=2)
            
            p = doc.add_paragraph()
            p.add_run("Decision Question: ").bold = True
            p.add_run(card.decision_question)
            
            # Status Section
            p2 = doc.add_paragraph()
            p2.add_run(f"Status: {s.status}  ").bold = True
            p2.add_run(f"(Priority Score: {s.total_priority:.2f})")

            # Evidence
            doc.add_heading("Key Evidence", level=3)
            if s.key_evidence:
                for ev in s.key_evidence:
                    doc.add_paragraph(ev, style='List Bullet')
            else:
                doc.add_paragraph("No critical triggers found.")

            # Recommendation
            if s.recommendation_draft:
                doc.add_heading("Recommendation Draft", level=3)
                rec = s.recommendation_draft
                
                tbl = doc.add_table(rows=4, cols=2)
                tbl.style = 'Table Grid'
                
                tbl.cell(0,0).text = "Action"
                tbl.cell(0,1).text = rec.action
                
                tbl.cell(1,0).text = "Risks"
                tbl.cell(1,1).text = str(rec.risks)
                
                tbl.cell(2,0).text = "Success Metrics"
                tbl.cell(2,1).text = str(rec.success_metrics)
                
                tbl.cell(3,0).text = "Preconditions"
                tbl.cell(3,1).text = str(rec.preconditions)

            # Human Override Log
            if s.human_override_reason:
                doc.add_heading("Audit Log", level=3)
                p_audit = doc.add_paragraph()
                p_audit.add_run("Override Reason: ").bold = True
                p_audit.add_run(s.human_override_reason)
            
            doc.add_paragraph() # Spacer
            doc.add_paragraph("_" * 40)

        # 3. Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
